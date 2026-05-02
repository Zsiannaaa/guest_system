from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"C:\xampp\htdocs\guest_system\docs\University_Guest_System_Testing_Walkthrough.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(8.5)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, title, rows):
    h = doc.add_heading(title, level=2)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    headers = ["ID", "Role", "Scenario / Step", "Expected Result", "Result", "Notes"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [0.55, 0.9, 3.15, 2.45, 0.65, 1.35]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, "1F7A35")
        set_cell_text(cell, header, bold=True, color="FFFFFF")
        cell.width = Inches(widths[idx])
    set_repeat_table_header(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        values = [
            row[0],
            row[1],
            row[2],
            row[3],
            "",
            "",
        ]
        for idx, value in enumerate(values):
            cells[idx].width = Inches(widths[idx])
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "EAF5E4")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 122, 53)
    r.font.size = Pt(10)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        run.font.size = Pt(9)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.add_run(item)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.55)
    sec.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.color.rgb = RGBColor(0, 95, 45)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(0, 95, 45)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("University Guest Monitoring and Visitor Management System")
    run.bold = True
    run.font.size = Pt(21)
    run.font.color.rgb = RGBColor(0, 95, 45)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("End-to-End Walkthrough Simulation Checklist")
    r.bold = True
    r.font.size = Pt(15)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("St. Paul University Dumaguete | PHP + MySQL | XAMPP Localhost | Prepared for system test drive").italic = True

    add_note(
        doc,
        "How to use this checklist",
        "Run the scenarios in order using a clean test day. Mark Result as Pass, Fail, or Blocked. Use Notes for screenshots, record IDs, reference numbers, and bugs found.",
    )

    doc.add_heading("Test Environment Setup", level=1)
    add_bullets(
        doc,
        [
            "Confirm Apache and MySQL are running in XAMPP.",
            "Open the system at http://localhost/guest_system/.",
            "Confirm database guest_system is imported and migrations are applied.",
            "Use test-only names, phone numbers, plate numbers, and organizations.",
            "Before destructive tests such as deleting guests, use records created specifically for this walkthrough.",
        ],
    )

    account_rows = [
        ("A1", "Admin", "Login as admin.", "Admin dashboard opens with summary cards, reports, users, offices, audit logs, and Guest House menu access."),
        ("A2", "Guard", "Login as guard1.", "Guard dashboard opens with pending arrivals, walk-in tools, check-in, checkout, and vehicle handling."),
        ("A3", "Office", "Login as one office staff account such as staff_reg.", "Office dashboard shows incoming visitors, currently serving guests, lookup, and office history."),
        ("A4", "Guest House", "Login as gh_staff.", "Guest House dashboard opens with expected guests, current occupants, rooms, bookings, and reports."),
        ("A5", "All", "Logout from each role and press browser Back.", "Protected dashboard pages must not reopen without authentication."),
    ]
    add_table(doc, "1. Login, Roles, and Session Checks", account_rows)

    doc.add_page_break()
    public_guest_rows = [
        ("P1", "Guest", "Open public pre-registration and submit a normal guest with full name, organization, purpose, date today/future, and one destination office.", "System creates a pending visit, reference number, and QR token."),
        ("P2", "Guest", "Repeat pre-registration with vehicle enabled: type, plate, color, model, university sticker/pass checked, optional sticker/pass number.", "Vehicle entry is saved with has vehicle, plate, and sticker/pass fields."),
        ("P3", "Guest", "Try a past visit date.", "Form rejects the date and shows a validation message."),
        ("P4", "Guest", "Open Check Status using the reference number.", "Only basic status is shown; private guest details are not exposed."),
        ("P5", "Guest", "Open Check Status using the QR token.", "Full visit details appear for QR token lookup."),
    ]
    add_table(doc, "2. Public Normal Guest Pre-Registration", public_guest_rows)

    guard_rows = [
        ("G1", "Guard", "Open pending check-in page and search by reference number.", "Pending visit is found with destination information."),
        ("G2", "Guard", "Open the same pending visit by QR token.", "Same visit is found and ready for check-in."),
        ("G3", "Guard", "Check in the pre-registered guest after verifying ID.", "Visit status becomes checked_in and activity log records check_in."),
        ("G4", "Guard", "Register a walk-in guest without vehicle.", "New guest profile and checked-in walk-in visit are created."),
        ("G5", "Guard", "Register a walk-in guest with vehicle and university sticker/pass checked.", "Vehicle details, sticker/pass flag, optional sticker/pass number, and driver are saved."),
        ("G6", "Guard", "Use saved guest record to create another check-in.", "Saved guest data is reused; a new visit reference is created."),
        ("G7", "Guard", "Try to check in a restricted guest.", "System blocks the check-in and shows restricted guest warning."),
        ("G8", "Guard", "Check out an active guest and confirm returned passes/badges if applicable.", "Visit status becomes checked_out and activity log records check_out."),
    ]
    add_table(doc, "3. Guard and Reception Workflow", guard_rows)

    doc.add_page_break()
    office_rows = [
        ("O1", "Office", "Login as office staff and review dashboard incoming visitors.", "Only guests routed to the office appear."),
        ("O2", "Office", "Confirm arrival for a routed checked-in guest.", "Destination status changes from pending to arrived."),
        ("O3", "Office", "Start service or mark in service if available.", "Destination status updates without affecting gate check-in status."),
        ("O4", "Office", "Complete the office destination.", "Destination status becomes completed and appears in office history."),
        ("O5", "Office", "Use Receive Visitor lookup for a guest not routed to this office.", "Office can add itself as an unplanned destination under the active visit."),
        ("O6", "Office", "Search by reference, QR token, and guest name.", "Checked-in visitors are searchable; checked-out visitors should not be received."),
    ]
    add_table(doc, "4. Office Staff Workflow", office_rows)

    admin_rows = [
        ("AD1", "Admin", "Create a new office.", "Office appears in destination selectors and office management list."),
        ("AD2", "Admin", "Edit the office name/location/status.", "Changes persist and inactive offices are hidden from active selectors."),
        ("AD3", "Admin", "Create a new user for guard, office staff, or Guest House staff.", "User can log in with assigned role and sees the correct dashboard."),
        ("AD4", "Admin", "Edit user details and reset password.", "Updated user information is saved; new password works."),
        ("AD5", "Admin", "Open guest directory, edit personal information.", "Guest profile changes persist without changing visit history."),
        ("AD6", "Admin", "Restrict/ban a guest with a reason.", "Guest is flagged as restricted and blocked from future check-in."),
        ("AD7", "Admin", "Lift restriction.", "Guest can be processed again after the flag is removed."),
        ("AD8", "Admin", "Attempt to delete a guest with visit or Guest House history.", "System blocks delete to protect history."),
        ("AD9", "Admin", "Delete a newly created guest with no history.", "Guest record is removed successfully."),
    ]
    add_table(doc, "5. Admin, Guest Directory, and Restriction Tests", admin_rows)

    doc.add_page_break()
    report_rows = [
        ("R1", "Admin", "Open main reports page with a date range.", "Summary cards, per-day visits, office breakdown, status breakdown, and guest log render."),
        ("R2", "Admin", "Export reports to CSV.", "CSV includes guest, office, status, vehicle, sticker/pass, and timing fields."),
        ("R3", "Admin", "Export one guest profile.", "CSV includes guest details, visit history, vehicle records, and sticker/pass fields."),
        ("R4", "Admin", "Open audit logs.", "Login, registration, check-in, checkout, restriction, office, and Guest House actions are traceable."),
    ]
    add_table(doc, "6. Reports, Exports, and Audit Logs", report_rows)

    gh_rows = [
        ("GH1", "GH Staff", "Create room types: Single, Double, Suite, Dormitory.", "Room types save and can be selected by rooms."),
        ("GH2", "GH Staff", "Create or edit rooms with capacity and status.", "Rooms appear in booking forms and availability views."),
        ("GH3", "GH Staff", "Create expected Guest House booking with guest name, organization, purpose, dates, room, and number of guests.", "Booking is reserved and guest profile is created or reused."),
        ("GH4", "GH Staff", "Try booking with past dates.", "System rejects past dates."),
        ("GH5", "GH Staff", "Try booking a room beyond capacity.", "System rejects over-capacity booking."),
        ("GH6", "GH Staff", "Try overlapping reservation for same room and dates.", "System blocks overlap."),
        ("GH7", "GH Staff", "Mark reserved booking as checked in.", "Booking becomes checked_in and room becomes occupied."),
        ("GH8", "GH Staff", "Review current occupants.", "Checked-in Guest House guest appears with room and stay details."),
        ("GH9", "GH Staff", "Generate linked campus visit from booking if guest will visit an office.", "A guest_visits record is created and linked to booking."),
        ("GH10", "GH Staff", "Check out Guest House occupant.", "Booking becomes checked_out and room returns to available when no active occupant remains."),
        ("GH11", "GH Staff", "Cancel a reserved booking.", "Booking becomes cancelled and no longer appears as active occupant."),
        ("GH12", "Admin/GH", "Open Guest House reports.", "Reports show expected guests, occupants, room usage, booking statuses, and date filters."),
    ]
    add_table(doc, "7. Guest House Accommodation Module", gh_rows)

    doc.add_page_break()
    security_rows = [
        ("S1", "Guest", "Open admin, guard, office, and Guest House pages while logged out.", "System redirects to login."),
        ("S2", "Office", "Try to open another office's destination handling URL.", "Access is denied unless the visit belongs to that office or user is admin."),
        ("S3", "All", "Submit forms twice or refresh after POST.", "System should avoid duplicate destructive actions via redirects and CSRF."),
        ("S4", "All", "Submit forms without CSRF token using browser tools or copied URL.", "System rejects request."),
        ("S5", "All", "Enter script tags in name, purpose, notes, and organization fields.", "Output is escaped; no script runs."),
        ("S6", "All", "Search using partial names, exact references, and QR tokens.", "Search works without exposing restricted or private data incorrectly."),
    ]
    add_table(doc, "8. Security and Edge Case Regression", security_rows)

    signoff_rows = [
        ("SO1", "Tester", "Record references for every test guest and booking.", "Traceability is available for cleanup and bug reports."),
        ("SO2", "Tester", "Clean up only test records that are safe to delete.", "Production-like history is preserved."),
        ("SO3", "Tester", "List failed cases with screenshot, URL, role, input values, and expected vs actual result.", "Bug report is actionable."),
        ("SO4", "Team", "Repeat failed cases after fixes.", "Regression result is documented."),
    ]
    add_table(doc, "9. Test Sign-Off and Cleanup", signoff_rows)

    doc.add_heading("Suggested Test Data", level=1)
    add_bullets(
        doc,
        [
            "Normal guest: Juan QA Visitor, organization QA Individual, purpose Registrar test.",
            "Vehicle guest: Maria Vehicle Test, plate QA 1234, sticker/pass SPU-TEST-001.",
            "Restricted guest: Restricted QA Visitor, reason System walkthrough restriction test.",
            "Guest House guest: VIP QA Visitor, organization Partner University, purpose Overnight official visit.",
            "Room tests: GH-TEST-101 Single, GH-TEST-201 Suite, and one maintenance room.",
        ],
    )

    doc.add_heading("Final Acceptance Criteria", level=1)
    add_bullets(
        doc,
        [
            "All roles can complete their core workflows without fatal errors.",
            "Normal guest visits, vehicles, sticker/pass data, destinations, and checkout are traceable.",
            "Restricted guests are blocked from check-in and visibly marked.",
            "Guest House bookings remain separate from regular campus visits but can link to a visit when needed.",
            "Reports and exports include the latest fields and match dashboard counts for the selected date range.",
            "Logout and browser Back behavior do not expose protected pages.",
        ],
    )

    section = doc.add_section(WD_SECTION.CONTINUOUS)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    doc.core_properties.title = "University Guest System Testing Walkthrough"
    doc.core_properties.subject = "End-to-end QA checklist for normal guest, office, guard, admin, and Guest House workflows"
    doc.core_properties.author = "Codex"
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
    print(OUTPUT)
